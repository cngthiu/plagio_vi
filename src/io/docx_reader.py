# =========================
# file: src/io/docx_reader.py
# =========================
from dataclasses import dataclass, field
from typing import List, Tuple, Optional
from docx import Document

@dataclass
class TableCell:
    text: str
    table_id: int
    row: int
    col: int
    start: int
    end: int

@dataclass
class DocumentStruct:
    paragraphs: List[str] = field(default_factory=list)
    tables: List[List[List[str]]] = field(default_factory=list)  # table -> rows -> cells
    linearized_text: str = ""
    mapping_cells: List[TableCell] = field(default_factory=list)

    def apply_ignores(self, headers_to_ignore: List[str]):
        # Why: giảm false positive từ mục chuẩn.
        if not headers_to_ignore:
            return
        lowered = [p.lower() for p in self.paragraphs]
        keep = []
        for p in lowered:
            if any(h.lower() in p for h in headers_to_ignore):
                continue
            keep.append(p)
        self.paragraphs = keep

    def linearize(self):
        parts = []
        offset = 0
        for p in self.paragraphs:
            parts.append(p)
            offset += len(p) + 1
        for tid, table in enumerate(self.tables):
            for r, row in enumerate(table):
                for c, cell in enumerate(row):
                    start = offset
                    parts.append(cell)
                    offset += len(cell) + 1
                    self.mapping_cells.append(TableCell(text=cell, table_id=tid, row=r, col=c, start=start, end=offset))
        self.linearized_text = "\n".join(parts)

def read_docx_with_tables(path: str, keep_tables: bool = True) -> DocumentStruct:
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    tables = []
    if keep_tables:
        for t in doc.tables:
            rows = []
            for row in t.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            tables.append(rows)
    return DocumentStruct(paragraphs=paragraphs, tables=tables)